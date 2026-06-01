class DocumentLoader:
    def __init__(self, file_paths):
        self.file_paths = file_paths

    def load_documents(self):
        from langchain.document import Document
        import os

        documents = []
        for file_path in self.file_paths:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            if file_path.endswith('.pdf'):
                documents.extend(self._load_pdf(file_path))
            elif file_path.endswith('.txt'):
                documents.append(self._load_txt(file_path))
            elif file_path.endswith('.docx'):
                documents.append(self._load_docx(file_path))
            else:
                raise ValueError(f"Unsupported file type: {file_path}")

        return documents

    def _load_pdf(self, file_path):
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        return [Document(page.extract_text()) for page in reader.pages]

    def _load_txt(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return [Document(content)]

    def _load_docx(self, file_path):
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        content = '\n'.join([para.text for para in doc.paragraphs])
        return [Document(content)]